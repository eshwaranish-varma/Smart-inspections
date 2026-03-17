"""FDA 483 & EIR document generation.

The 483 is built as ONE continuous table that spans all pages:
  - Header rows 0-5 are marked with <w:tblHeader/> so Word repeats them
    automatically on every printed/viewed page.
  - Each observation is a separate row that flows naturally across pages.
  - The footer (signature + form info) uses <w:cantSplit/> to stay together.
"""
from __future__ import annotations

import io
import logging
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Emu, Inches, Pt, Twips
from docx.table import Table

from app.config import settings
from app.schemas.inspection import (
    DraftObservation,
    EIRNarrative,
    InspectionMetadata,
)

logger = logging.getLogger(__name__)

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(10)
LABEL_SIZE = Pt(9)
HEADER_TITLE_SIZE = Pt(12)

TEMPLATE_NAME = "FDA_483a_Word_Template.docx"


def _find_template() -> Path:
    candidates = [
        settings.data_path / TEMPLATE_NAME,
        Path(__file__).resolve().parents[3] / "data" / TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(
        f"{TEMPLATE_NAME} not found in {[str(c) for c in candidates]}"
    )


# ── low-level XML helpers ──────────────────────────────────────────────


def _mark_row_as_header(row):
    """Add <w:tblHeader/> to a table row so it repeats on every page."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def _mark_row_cant_split(row):
    """Add <w:cantSplit/> so the row won't break across pages."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _set_cell_borders(cell, **kwargs):
    """Set borders on a table cell.  kwargs: top, bottom, left, right, each
    a dict with val, sz, color (all optional, defaults to single/4/000000).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge, props in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), props.get("val", "single"))
        el.set(qn("w:sz"), str(props.get("sz", 4)))
        el.set(qn("w:color"), props.get("color", "000000"))
        el.set(qn("w:space"), "0")
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        borders.append(el)


def _set_cell_width(cell, width_twips: int):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_vertical_alignment(cell, val="top"):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def _add_paragraph_bottom_border(paragraph):
    """Add a thin bottom border to a paragraph (acts as a divider line)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    old = pBdr.find(qn("w:bottom"))
    if old is not None:
        pBdr.remove(old)
    pBdr.append(bottom)


def _run(paragraph, text: str, bold=False, size=None, font_name=FONT_NAME):
    """Append a run to a paragraph with standard formatting."""
    r = paragraph.add_run(text)
    r.font.name = font_name
    r.font.size = size or FONT_SIZE
    r.bold = bold
    return r


def _label_value_block(cell, label: str, value: str):
    """Write a bold label on one line and the value on the next inside *cell*."""
    p = cell.paragraphs[0]
    p.clear()
    _run(p, label, bold=True, size=LABEL_SIZE)
    p2 = cell.add_paragraph()
    _run(p2, value, size=FONT_SIZE)


def _merge_row_cells(table, row_idx):
    """Merge all cells in a row into one spanning cell (horizontal merge)."""
    row = table.rows[row_idx]
    cells = row.cells
    if len(cells) <= 1:
        return cells[0]
    a = cells[0]
    for c in cells[1:]:
        a.merge(c)
    return a


# ── table-level border helper ──────────────────────────────────────────


def _set_table_borders(table, sz_outer="8", sz_inner="4", color="000000"):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz_outer)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        borders.append(el)
    for edge in ("insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz_inner)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        borders.append(el)


# ── main 483 builder ───────────────────────────────────────────────────

NUM_COLS = 3  # widest row has 3 columns; narrower rows merge cells


class DocumentService:

    @staticmethod
    def generate_483_document(
        form_data: InspectionMetadata,
        observations: list[DraftObservation],
    ) -> bytes:
        """Build a Form FDA 483 as ONE continuous table.

        Header rows (0-5) auto-repeat on every page via w:tblHeader.
        Each observation is its own row — Word flows them across pages.
        Footer rows stay together via w:cantSplit.
        """
        doc = Document()

        # ── page setup ──
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.375)
        section.right_margin = Inches(0.375)
        section.top_margin = Inches(0.375)
        section.bottom_margin = Inches(0.375)

        # Default style
        style = doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        usable_twips = 11160  # ~7.75 inches in twentieths of a point
        col_w = [3720, 3720, 3720]  # equal thirds

        table = doc.add_table(rows=0, cols=NUM_COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set table to 100% width
        tblPr = table._tbl.tblPr
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        _set_table_borders(table)

        # ────────────────────────────────────────────────────────────────
        # HEADER ROW 0: Title (merged across all cols)
        # ────────────────────────────────────────────────────────────────
        row0 = table.add_row()
        _mark_row_as_header(row0)
        c0 = _merge_row_cells(table, 0)
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "DEPARTMENT OF HEALTH AND HUMAN SERVICES", bold=True, size=HEADER_TITLE_SIZE)
        p2 = c0.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, "Food and Drug Administration", bold=True, size=HEADER_TITLE_SIZE)

        # ────────────────────────────────────────────────────────────────
        # HEADER ROW 1: District Office (col 0+1 merged) | Dates+FEI (col 2)
        # ────────────────────────────────────────────────────────────────
        row1 = table.add_row()
        _mark_row_as_header(row1)
        # Merge first two cells for district office
        left_cell = row1.cells[0].merge(row1.cells[1])
        right_cell = row1.cells[2]

        _label_value_block(left_cell, "DISTRICT OFFICE ADDRESS AND PHONE NUMBER:", form_data.district_office)
        _set_cell_vertical_alignment(left_cell, "top")

        p_dates = right_cell.paragraphs[0]
        p_dates.clear()
        _run(p_dates, "DATE(S) OF INSPECTION:", bold=True, size=LABEL_SIZE)
        pd2 = right_cell.add_paragraph()
        _run(pd2, f"{form_data.inspection_start} to {form_data.inspection_end}", size=FONT_SIZE)
        pd3 = right_cell.add_paragraph()
        _run(pd3, "FEI NUMBER:", bold=True, size=LABEL_SIZE)
        pd4 = right_cell.add_paragraph()
        _run(pd4, form_data.fei_number, size=FONT_SIZE)
        _set_cell_vertical_alignment(right_cell, "top")

        # ────────────────────────────────────────────────────────────────
        # HEADER ROW 2: Name & Title (merged)
        # ────────────────────────────────────────────────────────────────
        row2 = table.add_row()
        _mark_row_as_header(row2)
        c2 = _merge_row_cells(table, 2)
        p = c2.paragraphs[0]
        _run(p, "NAME AND TITLE OF INDIVIDUAL TO WHOM REPORT IS ISSUED:", bold=True, size=LABEL_SIZE)
        p2 = c2.add_paragraph()
        issued = f"{form_data.report_issued_to_name}"
        if form_data.report_issued_to_title:
            issued += f", {form_data.report_issued_to_title}"
        _run(p2, issued, size=FONT_SIZE)

        # ────────────────────────────────────────────────────────────────
        # HEADER ROW 3: Firm Name (col 0+1 merged) | Street Address (col 2)
        # ────────────────────────────────────────────────────────────────
        row3 = table.add_row()
        _mark_row_as_header(row3)
        firm_cell = row3.cells[0].merge(row3.cells[1])
        addr_cell = row3.cells[2]
        _label_value_block(firm_cell, "FIRM NAME:", form_data.firm_name)
        _label_value_block(addr_cell, "STREET ADDRESS:", form_data.street_address)

        # ────────────────────────────────────────────────────────────────
        # HEADER ROW 4: City/State/ZIP | Country | Type  (3 columns)
        # ────────────────────────────────────────────────────────────────
        row4 = table.add_row()
        _mark_row_as_header(row4)
        city_str = f"{form_data.city}, {form_data.state} {form_data.zip_code}"
        _label_value_block(row4.cells[0], "CITY, STATE AND ZIP CODE:", city_str)
        _label_value_block(row4.cells[1], "COUNTRY:", form_data.country)
        est_type = (form_data.establishment_type or "").replace("_", " ").title()
        _label_value_block(row4.cells[2], "TYPE OF ESTABLISHMENT INSPECTED:", est_type)

        # ────────────────────────────────────────────────────────────────
        # HEADER ROW 5: "During an inspection…" (merged)
        # ────────────────────────────────────────────────────────────────
        row5 = table.add_row()
        _mark_row_as_header(row5)
        c5 = _merge_row_cells(table, 5)
        p = c5.paragraphs[0]
        _run(p, "During an inspection of your establishment, I (we) observed:", size=FONT_SIZE)

        # ────────────────────────────────────────────────────────────────
        # OBSERVATION ROWS: one row per observation (merged, free-flowing)
        # ────────────────────────────────────────────────────────────────
        for idx, obs in enumerate(observations, start=1):
            obs_row = table.add_row()
            obs_cell = _merge_row_cells(table, len(table.rows) - 1)

            # Bold observation number
            p_label = obs_cell.paragraphs[0]
            _run(p_label, f"Observation {idx}:", bold=True, size=FONT_SIZE)

            # Observation body text
            p_body = obs_cell.add_paragraph()
            _run(p_body, obs.observation_text, size=FONT_SIZE)

            # CFR citation
            if obs.cfr_citation:
                p_cfr = obs_cell.add_paragraph()
                _run(p_cfr, f"[{obs.cfr_citation}]", bold=True, size=FONT_SIZE)

            # Evidence list
            if obs.evidence_list:
                for ev in obs.evidence_list:
                    if ev.strip():
                        p_ev = obs_cell.add_paragraph()
                        _run(p_ev, f"  \u2022 {ev}", size=FONT_SIZE)

            # Horizontal divider after each observation
            divider = obs_cell.add_paragraph()
            _add_paragraph_bottom_border(divider)

        # If no observations, add an empty placeholder row
        if not observations:
            empty_row = table.add_row()
            _merge_row_cells(table, len(table.rows) - 1)

        # ────────────────────────────────────────────────────────────────
        # FOOTER ROW A: Signature block (3 columns) — cantSplit
        # ────────────────────────────────────────────────────────────────
        sig_row = table.add_row()
        _mark_row_cant_split(sig_row)

        investigators = form_data.investigators or []
        inv_names = ", ".join(inv.get("name", "") for inv in investigators)
        inv_titles = "\n".join(
            f"{inv.get('name', '')}, {inv.get('title', '')}" for inv in investigators
        )

        _label_value_block(sig_row.cells[0], "EMPLOYEE(S) SIGNATURE:", inv_names)
        _label_value_block(sig_row.cells[1], "EMPLOYEE(S) NAME AND TITLE (Print or Type):", inv_titles)
        _label_value_block(sig_row.cells[2], "DATE ISSUED:", datetime.now().strftime("%m/%d/%Y"))

        # ────────────────────────────────────────────────────────────────
        # FOOTER ROW B: Form ID + Website + Page (merged) — cantSplit
        # ────────────────────────────────────────────────────────────────
        form_row = table.add_row()
        _mark_row_cant_split(form_row)
        form_cell = _merge_row_cells(table, len(table.rows) - 1)
        p = form_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(
            p,
            "FORM FDA 483 (2/16)       "
            "Website: www.fda.gov/oc/industry",
            bold=True,
            size=LABEL_SIZE,
        )

        # ── serialize ──
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ── EIR generation (unchanged) ─────────────────────────────────────

    @staticmethod
    def generate_eir_document(
        narrative: EIRNarrative, form_data: InspectionMetadata
    ) -> bytes:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("ESTABLISHMENT INSPECTION REPORT")
        r.bold = True
        r.font.size = Pt(14)

        doc.add_paragraph()
        for label, value in [
            ("Firm:", form_data.firm_name),
            ("FEI:", form_data.fei_number),
            ("Address:", f"{form_data.street_address}, {form_data.city}, {form_data.state} {form_data.zip_code}"),
            ("Inspection Dates:", f"{form_data.inspection_start} to {form_data.inspection_end}"),
            ("District Office:", form_data.district_office),
            ("Type:", (form_data.establishment_type or "").replace("_", " ").title()),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(label + " ")
            r.bold = True
            p.add_run(str(value))

        doc.add_page_break()
        for heading, content in [
            ("1. Background and Scope", narrative.background_scope),
            ("2. Observations Summary", narrative.observations_summary),
            ("3. Evidence Descriptions", narrative.evidence_descriptions),
            ("4. Chronological Account", narrative.chronological_account),
        ]:
            h = doc.add_heading(heading, level=2)
            if h.runs:
                h.runs[0].font.size = Pt(12)
            p = doc.add_paragraph()
            p.add_run(content or "(To be completed)").font.size = Pt(11)
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    @staticmethod
    def generate_both(
        form_data: InspectionMetadata,
        observations: list[DraftObservation],
        narrative: EIRNarrative,
    ) -> bytes:
        doc_483 = DocumentService.generate_483_document(form_data, observations)
        doc_eir = DocumentService.generate_eir_document(narrative, form_data)
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("FDA_483.docx", doc_483)
            zf.writestr("EIR_Narrative.docx", doc_eir)
        buf.seek(0)
        return buf.read()
